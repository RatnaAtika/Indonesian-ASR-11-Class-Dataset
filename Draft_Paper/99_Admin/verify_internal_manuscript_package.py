#!/usr/bin/env python3
"""Audit the NSS-ID internal manuscript package without authorizing submission."""
from __future__ import annotations

import argparse, csv, hashlib, json, re, sys, zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PIL import Image

ROOT=Path(__file__).resolve().parents[2]
DEFAULT_PACKAGE=ROOT/"Draft_Paper/05_Submission_Package"
DEFAULT_MD=ROOT/"Draft_Paper/03_Review/11_INTERNAL_PACKAGE_VERIFICATION_REPORT.md"
DEFAULT_JSON=ROOT/"Draft_Paper/03_Review/11_INTERNAL_PACKAGE_VERIFICATION_REPORT.json"
DRAFT=ROOT/"Draft_Paper/04_Revised_Draft"
SOURCE_DOCX_SHA256="17214b820dc3b70277541eeba1ca070de1cd2bd538e11ac66896c5957092bd0c"


def sha(path: Path) -> str: return hashlib.sha256(path.read_bytes()).hexdigest()

def rows(path: Path) -> list[dict[str,str]]:
    with path.open(newline="",encoding="utf-8-sig") as f: return list(csv.DictReader(f))

def docx_text(doc: Document) -> str:
    parts=[p.text for p in doc.paragraphs]
    for table in doc.tables:
        parts.extend(cell.text for row in table.rows for cell in row.cells)
    return "\n".join(parts)

def zip_text(path: Path) -> str:
    with zipfile.ZipFile(path) as z:
        return "\n".join(z.read(n).decode("utf-8","ignore") for n in z.namelist() if n.endswith((".xml",".rels")))

def package_text(package: Path) -> str:
    chunks=[]
    for path in sorted(package.rglob("*")):
        if path.is_file() and path.suffix.lower() in {".md",".csv",".json",".svg",".txt"}: chunks.append(path.read_text(encoding="utf-8-sig",errors="ignore"))
    chunks.append(zip_text(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"))
    chunks.append(zip_text(package/"NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx"))
    return "\n".join(chunks)


def check_source_custody(manifest: dict, source: Path) -> None:
    assert manifest["source_docx_sha256"]==SOURCE_DOCX_SHA256
    if source.exists(): assert sha(source)==SOURCE_DOCX_SHA256


def check_manifest(package: Path) -> str:
    manifest=json.loads((package/"PACKAGE_MANIFEST.json").read_text(encoding="utf-8"))
    assert manifest["status"]=="internal_not_for_submission"
    assert manifest["submission_authorized"] is False and manifest["public_release_authorized"] is False
    assert manifest["gate_status"]=={"G0":"NO-GO","G1":"NO-GO","G2":"NO-GO","G3":"NO-GO","G4":"NO-GO","G5":"NO-GO","G6":"UNASSESSED"}
    recorded={item["path"]:item for item in manifest["files"]}
    actual={p.relative_to(package).as_posix() for p in package.rglob("*") if p.is_file() and p.name!="PACKAGE_MANIFEST.json"}
    assert set(recorded)==actual,(set(recorded)^actual)
    for rel,item in recorded.items():
        path=package/rel; assert path.stat().st_size==item["bytes"]; assert sha(path)==item["sha256"],rel
    generator=ROOT/manifest["generator"]; assert sha(generator)==manifest["generator_sha256"]
    source=ROOT/"Draft_Paper/00_Source/Draft Jurnal Data In Brief NSS-ID_ver3.docx"
    check_source_custody(manifest,source)
    return f"{len(recorded)} listed files match SHA-256 and byte counts; G0–G5 NO-GO and G6 unassessed."


def check_material_gaps(package: Path) -> str:
    pattern=r"\[MATERIAL GAP:[^\]]+\]"
    canonical=set(re.findall(pattern,(DRAFT/"03_MATERIAL_GAP_PLACEHOLDERS.md").read_text(encoding="utf-8")))
    assert len(canonical)==33
    manuscript=(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md").read_text(encoding="utf-8")
    doc=Document(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx")
    for name,text in (("Markdown",manuscript),("DOCX",docx_text(doc))):
        found=re.findall(pattern,text); assert set(found)==canonical,name; assert len(found)>=33,name
    questionnaire=(package/"evidence/AUTHOR_METHODS_EVIDENCE_QUESTIONNAIRE.md").read_text(encoding="utf-8")
    assert canonical.issubset(set(re.findall(pattern,questionnaire)))
    matrix=rows(package/"evidence/METHODS_EVIDENCE_MATRIX.csv")
    assert len(matrix)>=60 and {row["classification"] for row in matrix}<={"OBSERVED","INFERRED","CONFLICTED","MISSING"}
    assert {"OBSERVED","INFERRED","CONFLICTED","MISSING"}.issubset({row["classification"] for row in matrix})
    flow_path=package/"evidence/CLAIM_EVIDENCE_FLOW.csv"; flow=rows(flow_path)
    assert len(flow)>=60 and all(None not in row for row in flow)
    flow_tokens=set(re.findall(r"\[MATERIAL GAP(?::[^]]+)?\]",flow_path.read_text(encoding="utf-8"))); assert "[MATERIAL GAP]" not in flow_tokens; assert flow_tokens.issubset(canonical),flow_tokens-canonical
    return "All 33 canonical tokens occur in Markdown, DOCX, and the author questionnaire; the methods matrix and structurally valid claim flow use controlled evidence and canonical closure language."


def check_quantitative(package: Path) -> str:
    t3=rows(package/"tables/Table_3_release_target_category_composition.csv")
    assert len(t3)==11 and sum(int(r["files"]) for r in t3)==104500
    assert sum(int(r["synthetic_files"]) for r in t3)==132
    assert sum(int(r["canonical_balanced_sentence_slots"]) for r in t3)==209
    assert abs(sum(float(r["duration_hours"]) for r in t3)-134.1763)<1e-8
    t4=rows(package/"tables/Table_4_release_target_split_source_composition.csv"); detail=[r for r in t4 if r["split"]!="Total"]; total=next(r for r in t4 if r["split"]=="Total")
    assert [int(r["files"]) for r in detail]==[73150,15675,15675]
    assert sum(int(r["files"]) for r in detail)==int(total["files"])==104500
    assert sum(int(r["human_recordings"]) for r in detail)==int(total["human_recordings"])==104368
    assert sum(int(r["synthetic_files"]) for r in detail)==int(total["synthetic_files"])==132
    assert abs(float(total["duration_hours"])-134.1762)<1e-8
    assert abs(sum(float(r["duration_hours"]) for r in detail)-134.1763)<1e-8
    registry=json.loads((package/"evidence/evidence_registry.json").read_text(encoding="utf-8"))
    assert registry["release_target_dataset"]["file_count"]==104500 and registry["benchmark_subset"]["file_count"]==102544
    assert registry["release_target_dataset"]["canonical_balanced_sentence_slots"]==209
    assert registry["benchmark_subset"]["local_source_validation"]["distinct_category_sentence_pairs"]==209
    return "Release target, split/source composition, canonical 209-slot design, frozen 102,544/209 scope, and disclosed rounded-hour difference reconcile."


def check_benchmark(package: Path) -> str:
    unified=rows(package/"evidence/unified_nine_model_metrics.csv"); table=rows(package/"tables/Table_S6_frozen_benchmark_validation.csv")
    assert len(unified)==len(table)==9
    assert {(r["reference_words"],r["reference_characters"],r["n_test_items"],r["normalizer_id"]) for r in unified}=={("135911","942599","15376","nssid_project_uniform_v1")}
    assert all(r["canonical_reference_match"]=="True" and r["audio_path_match"]=="True" for r in unified)
    source={r["model_family"]:r for r in unified}; assert [r["model_family"] for r in table]==sorted(source)
    for r in table:
        s=source[r["model_family"]]; assert r["wer_percent"]==f"{float(s['wer_percent']):.3f}"; assert r["cer_percent"]==f"{float(s['cer_percent']):.3f}"; assert r["parameters"]==s["parameters"]
    assert "rank" not in table[0] and "time" not in table[0]
    return "Nine rows use one 15,376-item normalizer and shared 135,911-word/942,599-character denominators; Supplementary Table S6 has no rank or timing."


def check_references(package: Path) -> str:
    refs=rows(package/"evidence/VERIFIED_REFERENCES.csv"); assert [int(r["ref_no"]) for r in refs]==list(range(1,17)); assert all(r["verification_status"].startswith("verified") for r in refs)
    manuscript=(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md").read_text(encoding="utf-8"); body=manuscript.split("# References",1)[0]
    order=[]
    for group in re.findall(r"\[([0-9,\-– ]+)\]",body):
        values=[]
        for part in group.split(","):
            part=part.strip()
            if re.fullmatch(r"\d+[-–]\d+",part):
                start,end=map(int,re.split(r"[-–]",part)); values.extend(range(start,end+1))
            elif part.isdigit(): values.append(int(part))
        for value in values:
            if value not in order: order.append(value)
    assert order==list(range(1,17)),order
    assert "10.15587/1729-4061.2026.350949" in manuscript
    return "Sixteen source-verified references are numbered and first cited in order; the related-article DOI is present."


def check_privacy_secrets_scope(package: Path) -> str:
    text=package_text(package)
    secret_or_path=[r"/mnt/c/Users/",r"/home/[^/\\s]+/",r"C:\\Users\\",r"\bhf_[A-Za-z0-9]{20,}\b",r"\bgh[pousr]_[A-Za-z0-9_]{20,}\b",r"\bsk-[A-Za-z0-9]{20,}\b",r"AKIA[0-9A-Z]{16}",r"BEGIN (?:RSA |EC |OPENSSH )?PRIVATE KEY",r"(?i)password\s*[:=]\s*[^\s,;]{6,}"]
    for pattern in secret_or_path: assert not re.search(pattern,text),pattern
    manuscript=(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md").read_text(encoding="utf-8")
    publication_text=manuscript+"\n"+docx_text(Document(package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"))+"\n"+"\n".join(p.read_text(encoding="utf-8-sig") for p in (package/"tables").glob("*.csv"))+"\n"+"\n".join(p.read_text(encoding="utf-8") for p in (package/"figures").glob("*.svg"))
    for pattern in (r"(?i)full public corpus",r"(?i)\banonymous\b",r"(?i)\bfully anonymized\b"): assert not re.search(pattern,publication_text),pattern
    allowed={"M":range(1,13),"F":range(1,9),"Ms":range(1,10),"Fs":range(1,10)}
    for prefix,number in re.findall(r"\b(Ms|Fs|M|F)(\d+)\b",manuscript): assert int(number) in allowed[prefix],(prefix,number)
    forbidden_ext={".wav",".flac",".mp3",".pt",".pth",".ckpt",".pkl",".npy",".npz",".env",".log"}
    assert not [(p.relative_to(package).as_posix(),p.suffix) for p in package.rglob("*") if p.is_file() and p.suffix.lower() in forbidden_ext]
    assert len(list(package.rglob("*.docx")))==1 and len(list(package.rglob("*.xlsx")))==1
    return "No targeted secret/path pattern, prohibited anonymity wording, out-of-range public ID, private-audio/model artifact, or forbidden file type was found."


def check_docx_xlsx(package: Path) -> str:
    docx=package/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"; doc=Document(docx); props=doc.core_properties
    assert (props.author or "")=="" and (props.last_modified_by or "")=="" and "NOT FOR SUBMISSION" in props.subject
    assert len(doc.tables)==6 and len(doc.inline_shapes)==3 and len(doc.sections)==1
    assert [row.cells[0].text for row in doc.tables[0].rows]==["Subject","Specific subject area","Type of data","Data collection","Data source location","Data accessibility","Related research article"]
    assert all("NOT FOR SUBMISSION" in " ".join(p.text for p in s.footer.paragraphs) for s in doc.sections)
    assert [p.text.strip() for p in doc.paragraphs if p.style and p.style.name=="Heading 1"]==["ARTICLE INFORMATION","SPECIFICATIONS TABLE","VALUE OF THE DATA","BACKGROUND","DATA DESCRIPTION","EXPERIMENTAL DESIGN, MATERIALS AND METHODS","LIMITATIONS","ETHICS STATEMENT","CRediT AUTHOR STATEMENT","ACKNOWLEDGEMENTS","DECLARATION OF COMPETING INTERESTS","REFERENCES"]
    with zipfile.ZipFile(docx) as z:
        assert z.testzip() is None and "word/vbaProject.bin" not in z.namelist(); assert all(i.date_time==(1980,1,1,0,0,0) for i in z.infolist())
        xml="\n".join(z.read(n).decode("utf-8","ignore") for n in z.namelist() if n.endswith((".xml",".rels")))
        assert "w:trackRevisions" not in xml and "documentProtection" not in xml and "TargetMode=\"External\"" not in xml and "comments.xml" not in z.namelist()
        media={sha256_bytes(z.read(n)) for n in z.namelist() if n.startswith("word/media/")}
    expected={sha(p) for p in (package/"figures").glob("*.png")}; header_hash="edb374c9b2e094a7b9e30adb7f988e7f684b6693e72ace19da1d5a9088db945a"; assert media==expected|{header_hash}
    xlsx=package/"NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx"; wb=load_workbook(xlsx)
    assert wb.sheetnames==["Specifications",*[f"Table {n}" for n in range(1,6)],"Table S6"]; assert (wb.properties.creator or "")=="" and (wb.properties.lastModifiedBy or "")==""
    assert all("NOT FOR SUBMISSION" in str(ws["A1"].value) for ws in wb.worksheets)
    with zipfile.ZipFile(xlsx) as z: assert z.testzip() is None and all(i.date_time==(1980,1,1,0,0,0) for i in z.infolist())
    return "Official-template DOCX has the exact 12-heading v.19 sequence, 6 tables (seven-row Specifications plus Tables 1–5), 3 body figures, retained header artwork, clean properties, and no comments/protection/external/macro/tracked-change parts; XLSX has 7 marked sheets including Supplementary Table S6."


def sha256_bytes(data: bytes) -> str: return hashlib.sha256(data).hexdigest()


def check_figures(package: Path) -> str:
    directory=package/"figures"; manifest=json.loads((directory/"figure_manifest.json").read_text(encoding="utf-8")); assert manifest["status"]=="internal_not_for_submission" and manifest["figure_4_status"]=="blocked_pending_sampling_provenance"
    for item in manifest["outputs"]: assert sha(directory/item["path"])==item["sha256"]
    for path in directory.glob("*.png"):
        with Image.open(path) as image:
            assert image.width>=3000 and image.height>=2000 and image.mode=="RGB"; dpi=image.info.get("dpi",(0,0)); assert dpi[0]>=590 and dpi[1]>=590; assert not image.getexif()
    svg="\n".join(p.read_text(encoding="utf-8") for p in directory.glob("*.svg")); assert "NOT FOR SUBMISSION" in svg and "/mnt/c/" not in svg and "/home/" not in svg
    return "Three 600-dpi PNG/SVG pairs match their figure manifest, contain no EXIF, remain visibly internal, and Figure 4 stays blocked."


def run_check(checks: list[dict], identifier: str, function, package: Path) -> None:
    try: checks.append({"id":identifier,"status":"PASS","detail":function(package)})
    except Exception as exc: checks.append({"id":identifier,"status":"FAIL","detail":f"{type(exc).__name__}: {exc}"})


def write_reports(checks: list[dict], md: Path, js: Path) -> bool:
    passed=all(c["status"]=="PASS" for c in checks); status="PASS_INTERNAL_ONLY" if passed else "FAIL"
    payload={"status":status,"submission_authorized":False,"public_release_authorized":False,"checks":checks,"residual_gate_status":{"G0":"NO-GO","G1":"NO-GO","G2":"NO-GO","G3":"NO-GO","G4":"NO-GO","G5":"NO-GO","G6":"UNASSESSED"}}
    js.parent.mkdir(parents=True,exist_ok=True); js.write_text(json.dumps(payload,indent=2)+"\n",encoding="utf-8")
    lines=["# Internal Package Verification Report","","> **NOT FOR SUBMISSION OR PUBLIC RELEASE**","",f"**Verification status:** `{status}`","","This status means only that the generated internal artifacts passed the listed mechanical checks. It does not close ethics, consent, rights, privacy, access, DOI, prior-publication, reproducibility, or author-authorization gates.",""]
    for c in checks: lines.extend([f"## {c['id']}: {c['status']}","",c["detail"],""])
    lines.extend(["## Gate decision","","- Internal rebuild: **GO**","- Journal submission: **NO-GO**","- Public release: **NO-GO**","- G0–G5: **NO-GO**","- G6: **UNASSESSED**",""])
    md.write_text("\n".join(lines),encoding="utf-8"); return passed


def main() -> None:
    parser=argparse.ArgumentParser(description=__doc__); parser.add_argument("--package-dir",type=Path,default=DEFAULT_PACKAGE); parser.add_argument("--report-md",type=Path,default=DEFAULT_MD); parser.add_argument("--report-json",type=Path,default=DEFAULT_JSON); args=parser.parse_args(); package=args.package_dir.resolve()
    checks=[]
    for identifier,function in (("package_manifest",check_manifest),("material_gaps",check_material_gaps),("quantitative_consistency",check_quantitative),("benchmark_comparability",check_benchmark),("citations",check_references),("privacy_secrets_file_scope",check_privacy_secrets_scope),("docx_xlsx_integrity",check_docx_xlsx),("figure_integrity",check_figures)): run_check(checks,identifier,function,package)
    passed=write_reports(checks,args.report_md.resolve(),args.report_json.resolve()); print(json.dumps({"status":"PASS_INTERNAL_ONLY" if passed else "FAIL","passed":sum(c["status"]=="PASS" for c in checks),"total":len(checks)},indent=2)); sys.exit(0 if passed else 1)


if __name__=="__main__": main()
