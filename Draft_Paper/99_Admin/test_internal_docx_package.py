#!/usr/bin/env python3
"""Integration tests for the deterministic internal DOCX/XLSX package builder."""

from __future__ import annotations

import hashlib
import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

ROOT = Path(__file__).resolve().parents[2]
SCRIPT = ROOT / "Draft_Paper" / "99_Admin" / "build_internal_docx_package.py"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


class InternalDocxPackageTests(unittest.TestCase):
    def build(self, output: Path) -> None:
        subprocess.run(
            [sys.executable, str(SCRIPT), "--output-dir", str(output)],
            cwd=ROOT,
            check=True,
            timeout=240,
        )

    def test_package_is_complete_sanitized_and_deterministic(self) -> None:
        with tempfile.TemporaryDirectory() as first_dir, tempfile.TemporaryDirectory() as second_dir:
            first = Path(first_dir)
            second = Path(second_dir)
            self.build(first)
            self.build(second)

            docx_path = first / "NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"
            xlsx_path = first / "NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx"
            manifest_path = first / "PACKAGE_MANIFEST.json"
            readme_path = first / "README_INTERNAL_NOT_FOR_SUBMISSION.md"
            self.assertTrue(docx_path.is_file())
            self.assertTrue(xlsx_path.is_file())
            self.assertTrue(manifest_path.is_file())
            self.assertTrue(readme_path.is_file())
            self.assertEqual(sha256(docx_path), sha256(second / docx_path.name))
            self.assertEqual(sha256(xlsx_path), sha256(second / xlsx_path.name))

            document = Document(docx_path)
            props = document.core_properties
            self.assertEqual("", props.author or "")
            self.assertEqual("", props.last_modified_by or "")
            self.assertIn("NOT FOR SUBMISSION", props.subject)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("INTERNAL WORKING DRAFT — NOT FOR SUBMISSION OR PUBLIC RELEASE", text)
            self.assertIn("104,500 PCM WAV files totaling 134.1762 h", text)
            self.assertIn("uniform diagnostic rescore", text.lower())
            self.assertIn("EXPERIMENTAL DESIGN, MATERIALS AND METHODS", text)
            self.assertNotIn("4. Experimental design, materials and methods", text)
            self.assertIn("110,000 source WAV files", text)
            self.assertIn("approximately 1.5037 m × 2.5027 m", text)
            self.assertIn("[MATERIAL GAP: ethics committee/determination, reference number, and date]", text)
            self.assertNotIn("/mnt/c/Users/", text)
            self.assertNotRegex(text, r"/home/[^/\s]+/")
            self.assertEqual(6, len(document.tables))
            self.assertEqual(
                [
                    "Subject",
                    "Specific subject area",
                    "Type of data",
                    "Data collection",
                    "Data source location",
                    "Data accessibility",
                    "Related research article",
                ],
                [row.cells[0].text for row in document.tables[0].rows],
            )
            self.assertEqual(3, len(document.inline_shapes))
            self.assertEqual(1, len(document.sections))
            for section in document.sections:
                footer_text = " ".join(paragraph.text for paragraph in section.footer.paragraphs)
                self.assertIn("NOT FOR SUBMISSION", footer_text)

            with zipfile.ZipFile(docx_path) as archive:
                names = archive.namelist()
                self.assertNotIn("word/vbaProject.bin", names)
                self.assertTrue(all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist()))
                xml_text = "\n".join(
                    archive.read(name).decode("utf-8", "ignore")
                    for name in names
                    if name.endswith(".xml") or name.endswith(".rels")
                )
                self.assertNotIn("/mnt/c/Users/", xml_text)
                self.assertNotRegex(xml_text, r"/home/[^/\s]+/")
                self.assertNotIn("w:trackRevisions", xml_text)
                self.assertNotIn("documentProtection", xml_text)
                self.assertNotIn("TargetMode=\"External\"", xml_text)
                self.assertNotIn("comments.xml", names)

            workbook = load_workbook(xlsx_path, read_only=False, data_only=False)
            self.assertEqual(
                ["Specifications", "Table 1", "Table 2", "Table 3", "Table 4", "Table 5", "Table S6"],
                workbook.sheetnames,
            )
            self.assertEqual("", workbook.properties.creator or "")
            self.assertEqual("", workbook.properties.lastModifiedBy or "")
            for worksheet in workbook.worksheets:
                self.assertIn("NOT FOR SUBMISSION", str(worksheet["A1"].value))
                self.assertGreater(worksheet.max_row, 5)
            with zipfile.ZipFile(xlsx_path) as archive:
                self.assertTrue(all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist()))

            self.assertEqual(7, len(list((first / "tables").glob("*.csv"))))
            self.assertEqual(7, len(list((first / "figures").iterdir())))
            self.assertGreaterEqual(len(list((first / "evidence").iterdir())), 21)
            for name in [
                "METHODS_EVIDENCE_MATRIX.csv",
                "AUTHOR_METHODS_EVIDENCE_QUESTIONNAIRE.md",
                "CURRENT_DIB_SPEECH_METHODS_EXPECTATIONS.md",
                "MANUSCRIPT_ARCHITECTURE.md",
                "CLAIM_EVIDENCE_FLOW.csv",
                "DISPLAY_AND_SUPPLEMENT_PLAN.md",
                "DATA_IN_BRIEF_TEMPLATE_MIGRATION_DESIGN.md",
                "OFFICIAL_TEMPLATE_INVENTORY.json",
            ]:
                self.assertTrue((first / "evidence" / name).is_file(), name)

            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual("internal_not_for_submission", manifest["status"])
            self.assertFalse(manifest["submission_authorized"])
            self.assertFalse(manifest["public_release_authorized"])
            recorded = {item["path"]: item for item in manifest["files"]}
            self.assertIn(docx_path.name, recorded)
            self.assertIn(xlsx_path.name, recorded)
            self.assertIn("NSS-ID_EVIDENCE_MASTER_INTERNAL.md", recorded)
            self.assertEqual(
                "5c02d5f9e0762e05f69c06d1d042ea800b6214427c82a78166863dfd17264190",
                manifest["official_template_sha256"],
            )
            self.assertEqual("v.19 (December 2024)", manifest["official_template_version"])
            for relative, item in recorded.items():
                path = first / relative
                self.assertTrue(path.is_file(), relative)
                self.assertEqual(sha256(path), item["sha256"])


if __name__ == "__main__":
    unittest.main()
