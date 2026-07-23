#!/usr/bin/env python3
"""Build a deterministic internal NSS-ID DOCX from Data in Brief template v.19."""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "Draft_Paper" / "data-in-brief-article-template.docx"
SOURCE = ROOT / "Draft_Paper" / "04_Revised_Draft" / "06_DATA_IN_BRIEF_TEMPLATE_MANUSCRIPT.md"
MASTER = ROOT / "Draft_Paper" / "04_Revised_Draft" / "04_INTERNAL_WORKING_MANUSCRIPT.md"
CAPTIONS = ROOT / "Draft_Paper" / "04_Revised_Draft" / "05_TABLE_CAPTIONS_AND_NOTES.md"
TABLE_DIR = ROOT / "Draft_Paper" / "04_Revised_Draft" / "tables"
FIGURE_DIR = ROOT / "Draft_Paper" / "04_Revised_Draft" / "figures"
DEFAULT_OUTPUT = ROOT / "Draft_Paper" / "05_Submission_Package" / "NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"
FIXED_DATE = datetime(2026, 7, 23)
TEMPLATE_SHA256 = "5c02d5f9e0762e05f69c06d1d042ea800b6214427c82a78166863dfd17264190"

TABLE_FILES = {
    1: "Table_1_package_inventory.csv",
    2: "Table_2_scope_bridge.csv",
    3: "Table_3_release_target_category_composition.csv",
    4: "Table_4_release_target_split_source_composition.csv",
    5: "Table_5_synthetic_repair_provenance.csv",
}
FIGURE_FILES = {
    1: "Figure_1_construction_package_flow.png",
    2: "Figure_2_release_target_duration_by_category.png",
    3: "Figure_3_release_target_split_source_composition.png",
}
DISPLAY_FIELDS = {
    1: ["component", "path_or_archive_member", "format", "rows_or_files", "package_state"],
    2: ["field", "release_target", "frozen_benchmark", "evidence_control"],
    3: [
        "category_english",
        "files",
        "duration_hours",
        "mean_duration_sec",
        "synthetic_files",
        "distinct_category_sentence_id_pairs",
        "sentence_id_note",
    ],
    4: [
        "split",
        "human_speakers",
        "files",
        "human_recordings",
        "synthetic_files",
        "duration_hours",
        "male_source_files",
        "female_source_files",
    ],
    5: ["dimension", "value", "files", "duration_sec", "percent_of_release_target", "note", "source_scope"],
}

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
NS = {"w": W_NS, "r": REL_NS, "ct": CT_NS}

REMOVED_PARTS = {
    "word/comments.xml",
    "word/commentsExtended.xml",
    "word/commentsIds.xml",
    "word/commentsExtensible.xml",
    "word/people.xml",
    "word/_rels/comments.xml.rels",
    "docProps/custom.xml",
}
REMOVED_RELATIONSHIP_SUFFIXES = {
    "comments",
    "commentsExtended",
    "commentsIds",
    "commentsExtensible",
    "person",
    "custom-properties",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def plain(text: str) -> str:
    def replace_link(match: re.Match[str]) -> str:
        label, target = match.group(1), match.group(2)
        return f"{label} ({target})" if target.startswith(("http://", "https://")) else label

    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", replace_link, text)
    return text.replace("**", "").replace("`", "").strip()


def clear_body(doc: Document) -> None:
    body = doc._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shade = properties.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        properties.append(shade)
    shade.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shade = properties.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        properties.append(shade)
    shade.set(qn("w:fill"), fill)


def repeat_header(row) -> None:
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(node)


def add_inline(paragraph, text: str) -> None:
    # Repository-navigation links remain clickable in Markdown but render as
    # their human-readable labels in the standalone DOCX.
    text = re.sub(r"\[([^]]+)\]\((?!https?://)[^)]+\)", r"\1", text)
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            if "MATERIAL GAP" in token:
                run.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def csv_rows(path: Path) -> tuple[list[str], list[dict[str, str]]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)
        return list(reader.fieldnames or []), rows


def pretty(field: str) -> str:
    names = {
        "path_or_archive_member": "Path / member",
        "rows_or_files": "Rows / files",
        "package_state": "Package state",
        "release_target": "Release target",
        "frozen_benchmark": "Frozen benchmark",
        "evidence_control": "Evidence control",
        "category_english": "Category",
        "duration_hours": "Duration (h)",
        "mean_duration_sec": "Mean duration (s)",
        "synthetic_files": "Synthetic",
        "distinct_category_sentence_id_pairs": "Category–ID pairs",
        "sentence_id_note": "ID note",
        "human_speakers": "Human public labels",
        "human_recordings": "Human recordings",
        "male_source_files": "Male-source files",
        "female_source_files": "Female-source files",
        "interpretation_note": "Scope note",
        "duration_sec": "Duration (s)",
        "percent_of_release_target": "% of release target",
        "source_scope": "Source / scope",
    }
    return names.get(field, field.replace("_", " ").title())


def table_captions() -> dict[int, str]:
    captions: dict[int, str] = {}
    current: int | None = None
    for line in CAPTIONS.read_text(encoding="utf-8").splitlines():
        match = re.match(r"## Table (\d+)\.", line)
        if match:
            current = int(match.group(1))
        elif current is not None and line.startswith("**Caption:**"):
            captions[current] = plain(line.split("**Caption:**", 1)[1])
            current = None
    if set(captions) != set(TABLE_FILES):
        raise ValueError(f"Unexpected main-table caption coverage: {sorted(captions)}")
    return captions


def figure_captions() -> dict[int, str]:
    captions: dict[int, str] = {}
    for line in MASTER.read_text(encoding="utf-8").splitlines():
        match = re.match(r"\*\*Figure (\d+)\. (.+?)\*\*\s*(.*)", line)
        if match:
            captions[int(match.group(1))] = f"Figure {match.group(1)}. {match.group(2)} {match.group(3)}".strip()
    return captions


def add_specifications_table(doc: Document) -> None:
    _, rows = csv_rows(TABLE_DIR / "Specifications_Table.csv")
    expected = [
        "Subject",
        "Specific subject area",
        "Type of data",
        "Data collection",
        "Data source location",
        "Data accessibility",
        "Related research article",
    ]
    if [row["item"] for row in rows] != expected:
        raise ValueError("Specifications_Table.csv does not match Data in Brief v.19 fixed rows")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, row in enumerate(rows):
        cells = table.add_row().cells
        cells[0].width = Inches(1.3)
        cells[1].width = Inches(5.15)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[0].text = row["item"]
        # The official v.19 limits apply to the complete rendered cell. Detailed
        # evidence status remains in the editable CSV and evidence registry.
        cells[1].text = row["description"]
        shade_cell(cells[0], "E7E6E6")
        if index % 2:
            shade_cell(cells[1], "F7F7F7")
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for paragraph in cells[1].paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(8)
                if "MATERIAL GAP" in run.text:
                    run.font.color.rgb = RGBColor(192, 0, 0)
    source = doc.add_paragraph()
    source_run = source.add_run("Editable source: tables/Specifications_Table.csv")
    source_run.italic = True
    source_run.font.size = Pt(8)


def add_data_table(doc: Document, number: int, caption: str) -> None:
    _, source_rows = csv_rows(TABLE_DIR / TABLE_FILES[number])
    fields = DISPLAY_FIELDS[number]
    caption_paragraph = doc.add_paragraph()
    caption_run = caption_paragraph.add_run(f"Table {number}. {caption}")
    caption_run.font.size = Pt(8)
    caption_run.bold = True
    table = doc.add_table(rows=1, cols=len(fields))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    repeat_header(table.rows[0])
    font_size = Pt(5.5 if len(fields) >= 7 else 6.5)
    for index, field in enumerate(fields):
        cell = table.rows[0].cells[index]
        cell.text = pretty(field)
        shade_cell(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = font_size
    for row_index, row in enumerate(source_rows, start=1):
        cells = table.add_row().cells
        for index, field in enumerate(fields):
            cells[index].text = plain(row.get(field, ""))
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index % 2 == 0:
                shade_cell(cells[index], "F4F8FB")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = font_size
    source = doc.add_paragraph()
    source_run = source.add_run(f"Editable full source: tables/{TABLE_FILES[number]}")
    source_run.italic = True
    source_run.font.size = Pt(8)


def add_figure(doc: Document, number: int, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(FIGURE_DIR / FIGURE_FILES[number]), width=Inches(6.25))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(plain(caption))
    caption_run.bold = True
    caption_run.font.size = Pt(8)


def set_internal_marking(doc: Document) -> None:
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(banner, "FCE4D6")
    run = banner.add_run("INTERNAL WORKING DRAFT — NOT FOR SUBMISSION OR PUBLIC RELEASE")
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.font.size = Pt(11)
    for section in doc.sections:
        for footer in (section.footer, section.even_page_footer, section.first_page_footer):
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("INTERNAL — NOT FOR SUBMISSION OR PUBLIC RELEASE")
            run.bold = True
            run.font.color.rgb = RGBColor(192, 0, 0)
            run.font.size = Pt(8)


def populate_document(doc: Document) -> None:
    clear_body(doc)
    set_internal_marking(doc)
    table_notes = table_captions()
    figure_notes = figure_captions()
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if not line or line.startswith("# INTERNAL WORKING DRAFT"):
            continue
        if line == "[SPECIFICATIONS_TABLE]":
            add_specifications_table(doc)
            continue
        table_match = re.fullmatch(r"\[TABLE (\d+)\]", line)
        if table_match:
            number = int(table_match.group(1))
            add_data_table(doc, number, table_notes[number])
            continue
        figure_match = re.fullmatch(r"\[FIGURE (\d+)\]", line)
        if figure_match:
            number = int(figure_match.group(1))
            add_figure(doc, number, figure_notes[number])
            continue
        heading = re.match(r"^##\s+(.+)$", line)
        if heading:
            doc.add_heading(plain(heading.group(1)), level=1)
            continue
        subheading = re.match(r"^###\s+(.+)$", line)
        if subheading:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run(plain(subheading.group(1)))
            run.bold = True
            run.font.size = Pt(11)
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Paragraph")
            add_inline(paragraph, "• " + line[2:])
            continue
        paragraph = doc.add_paragraph()
        if re.match(r"^\[\d+\]\s", line) or line.startswith("[dataset]"):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        add_inline(paragraph, line)


def sanitize_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "NSS-ID Data in Brief template-aligned internal working manuscript"
    props.subject = "INTERNAL — NOT FOR SUBMISSION OR PUBLIC RELEASE"
    props.author = ""
    props.last_modified_by = ""
    props.keywords = "internal; Data in Brief template v.19; not for submission"
    props.comments = "Built from the official v.19 template; unresolved material gaps remain."
    props.category = "Internal template-aligned data article"
    props.created = FIXED_DATE
    props.modified = FIXED_DATE
    props.revision = 1


def clean_xml_part(name: str, data: bytes) -> bytes:
    if name == "word/settings.xml":
        root = etree.fromstring(data)
        for node in root.xpath("//w:documentProtection", namespaces={"w": W_NS}):
            node.getparent().remove(node)
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if name == "word/document.xml":
        root = etree.fromstring(data)
        removable = [
            "//w:commentRangeStart",
            "//w:commentRangeEnd",
            "//w:commentReference",
            "//w:permStart",
            "//w:permEnd",
            "//w:ins",
            "//w:del",
        ]
        for expression in removable:
            for node in root.xpath(expression, namespaces={"w": W_NS}):
                node.getparent().remove(node)
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if name.endswith(".rels"):
        root = etree.fromstring(data)
        for relation in list(root):
            rel_type = (relation.get("Type") or "").rsplit("/", 1)[-1]
            target = relation.get("Target") or ""
            if relation.get("TargetMode") == "External" or rel_type in REMOVED_RELATIONSHIP_SUFFIXES:
                root.remove(relation)
            elif any(target.endswith(Path(part).name) for part in REMOVED_PARTS):
                root.remove(relation)
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if name == "[Content_Types].xml":
        root = etree.fromstring(data)
        removed_names = {"/" + part for part in REMOVED_PARTS}
        for node in list(root):
            if node.get("PartName") in removed_names:
                root.remove(node)
        return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    if name == "docProps/core.xml":
        data = re.sub(
            rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",
            rb"\g<1>2026-07-23T00:00:00Z\g<2>",
            data,
        )
    return data


def normalize_and_sanitize_zip(path: Path) -> None:
    with zipfile.ZipFile(path, "r") as source:
        entries: list[tuple[zipfile.ZipInfo, bytes]] = []
        for info in source.infolist():
            if info.filename in REMOVED_PARTS:
                continue
            entries.append((info, clean_xml_part(info.filename, source.read(info.filename))))
    with tempfile.NamedTemporaryFile(delete=False, dir=path.parent, suffix=path.suffix) as handle:
        temporary = Path(handle.name)
    try:
        with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
            for old, data in sorted(entries, key=lambda item: item[0].filename):
                info = zipfile.ZipInfo(old.filename, (1980, 1, 1, 0, 0, 0))
                info.compress_type = zipfile.ZIP_STORED if old.is_dir() else zipfile.ZIP_DEFLATED
                info.external_attr = old.external_attr
                info.create_system = 0
                target.writestr(info, data)
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def build(output: Path) -> None:
    if sha256(TEMPLATE) != TEMPLATE_SHA256:
        raise ValueError("Official template hash changed; re-audit before building")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    populate_document(doc)
    sanitize_core_properties(doc)
    doc.save(output)
    normalize_and_sanitize_zip(output)
    print(f"Built Data in Brief v.19 template-aligned internal DOCX: {output}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    build(args.output.resolve())


if __name__ == "__main__":
    main()
