#!/usr/bin/env python3
"""Build deterministic NSS-ID DOCX/XLSX artifacts for internal review only."""
from __future__ import annotations

import argparse, csv, hashlib, json, os, re, shutil, tempfile, zipfile
from datetime import datetime
from pathlib import Path
from urllib.parse import quote, unquote, urlsplit

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

import build_data_in_brief_template_docx as dib_template

ROOT = Path(__file__).resolve().parents[2]
DRAFT = ROOT / "Draft_Paper/04_Revised_Draft"
EVIDENCE = ROOT / "Draft_Paper/02_Evidence"
REVIEWS = ROOT / "Draft_Paper/03_Review"
DEFAULT_OUTPUT = ROOT / "Draft_Paper/05_Submission_Package"
EVIDENCE_MASTER = DRAFT / "04_INTERNAL_WORKING_MANUSCRIPT.md"
MANUSCRIPT = EVIDENCE_MASTER
TEMPLATE_MANUSCRIPT = DRAFT / "06_DATA_IN_BRIEF_TEMPLATE_MANUSCRIPT.md"
CAPTIONS = DRAFT / "05_TABLE_CAPTIONS_AND_NOTES.md"
TABLE_DIR, FIGURE_DIR = DRAFT / "tables", DRAFT / "figures"
FIXED_DATE = datetime(2026, 7, 22)
SOURCE_DOCX_SHA256 = "17214b820dc3b70277541eeba1ca070de1cd2bd538e11ac66896c5957092bd0c"
DOCX_NAME = "NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx"
XLSX_NAME = "NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx"
TABLE_FILES = {
  1:"Table_1_package_inventory.csv", 2:"Table_2_scope_bridge.csv",
  3:"Table_3_release_target_category_composition.csv",
  4:"Table_4_release_target_split_source_composition.csv",
  5:"Table_5_synthetic_repair_provenance.csv",
  6:"Table_S6_frozen_benchmark_validation.csv",
}
TABLE_TRIGGERS = {
  1:"Table 1 distinguishes artifacts", 2:"Table 2 is the mandatory bridge",
  3:"Figure 2 and Table 3 use", 4:"Table 4 separates",
  5:"Table 5 records", 6:"Supplementary Table S6 records",
}
FIGURE_FILES = {
  1:"Figure_1_construction_package_flow.png",
  2:"Figure_2_release_target_duration_by_category.png",
  3:"Figure_3_release_target_split_source_composition.png",
}
SHEETS = [
 ("Specifications","Specifications_Table.csv"),
 *[(f"Table {n}", TABLE_FILES[n]) for n in range(1,6)],
 ("Table S6", TABLE_FILES[6]),
]
EVIDENCE_FILES = [
 (EVIDENCE/"EVIDENCE_REGISTRY.md","EVIDENCE_REGISTRY.md"),
 (EVIDENCE/"evidence_registry.json","evidence_registry.json"),
 (EVIDENCE/"claim_evidence_matrix.csv","claim_evidence_matrix.csv"),
 (EVIDENCE/"VERIFIED_REFERENCES.csv","VERIFIED_REFERENCES.csv"),
 (EVIDENCE/"PRIOR_PUBLICATION_OVERLAP_ASSESSMENT.md","PRIOR_PUBLICATION_OVERLAP_ASSESSMENT.md"),
 (EVIDENCE/"JOURNAL_REQUIREMENTS_SNAPSHOT.md","JOURNAL_REQUIREMENTS_SNAPSHOT.md"),
 (EVIDENCE/"CURRENT_DIB_SPEECH_METHODS_EXPECTATIONS.md","CURRENT_DIB_SPEECH_METHODS_EXPECTATIONS.md"),
 (EVIDENCE/"METHODS_EVIDENCE_MATRIX.csv","METHODS_EVIDENCE_MATRIX.csv"),
 (EVIDENCE/"AUTHOR_METHODS_EVIDENCE_QUESTIONNAIRE.md","AUTHOR_METHODS_EVIDENCE_QUESTIONNAIRE.md"),
 (EVIDENCE/"unified_benchmark_rescore/BENCHMARK_SCORING_COMPARABILITY_AUDIT.md","BENCHMARK_SCORING_COMPARABILITY_AUDIT.md"),
 (EVIDENCE/"unified_benchmark_rescore/unified_nine_model_metrics.csv","unified_nine_model_metrics.csv"),
 (EVIDENCE/"unified_benchmark_rescore/unified_nine_model_metrics.json","unified_nine_model_metrics.json"),
 (DRAFT/"00_MANUSCRIPT_ARCHITECTURE.md","MANUSCRIPT_ARCHITECTURE.md"),
 (DRAFT/"01_CLAIM_EVIDENCE_FLOW.csv","CLAIM_EVIDENCE_FLOW.csv"),
 (ROOT/"Draft_Paper/99_Admin/DATA_IN_BRIEF_TEMPLATE_MIGRATION_DESIGN.md","DATA_IN_BRIEF_TEMPLATE_MIGRATION_DESIGN.md"),
 (EVIDENCE/"OFFICIAL_DATA_IN_BRIEF_TEMPLATE_INVENTORY.json","OFFICIAL_TEMPLATE_INVENTORY.json"),
 (DRAFT/"02_DISPLAY_AND_SUPPLEMENT_PLAN.md","DISPLAY_AND_SUPPLEMENT_PLAN.md"),
 (DRAFT/"03_MATERIAL_GAP_PLACEHOLDERS.md","MATERIAL_GAP_PLACEHOLDERS.md"),
 (REVIEWS/"07_CONSOLIDATED_GAP_ANALYSIS.md","CONSOLIDATED_GAP_ANALYSIS.md"),
 (REVIEWS/"08_REVIEWER_RISK_MATRIX.csv","REVIEWER_RISK_MATRIX.csv"),
 (REVIEWS/"09_PRIORITIZED_REVISION_ROADMAP.md","PRIORITIZED_REVISION_ROADMAP.md"),
 (REVIEWS/"10_PROJECT_RESUME_AND_READINESS.md","PROJECT_RESUME_AND_READINESS.md"),
]


def sha256(path: Path) -> str:
    h=hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda:f.read(1<<20),b""): h.update(chunk)
    return h.hexdigest()


def plain(text: str) -> str:
    def replace_link(match: re.Match[str]) -> str:
        label,target=match.group(1),match.group(2)
        return f"{label} ({target})" if target.startswith(("http://","https://")) else label
    text=re.sub(r"\[([^]]+)\]\(([^)]+)\)",replace_link,text)
    return re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)",r"\1",text.replace("**","").replace("`","" )).strip()


def add_inline(p, text: str) -> None:
    text=re.sub(r"\[([^]]+)\]\((?!https?://)[^)]+\)",r"\1",text)
    pattern=re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)"); pos=0
    for m in pattern.finditer(text):
        if m.start()>pos: p.add_run(text[pos:m.start()])
        token=m.group(0); run=p.add_run(token[2:-2] if token.startswith("**") else token[1:-1])
        if token.startswith("**"): run.bold=True
        elif token.startswith("`"):
            run.font.name="Consolas"; run.font.size=Pt(9)
            if "MATERIAL GAP" in token: run.bold=True; run.font.color.rgb=RGBColor(201,76,76)
        else: run.italic=True
        pos=m.end()
    if pos<len(text): p.add_run(text[pos:])


def shade(cell, fill: str) -> None:
    pr=cell._tc.get_or_add_tcPr(); node=pr.find(qn("w:shd"))
    if node is None: node=OxmlElement("w:shd"); pr.append(node)
    node.set(qn("w:fill"),fill)


def repeat_header(row) -> None:
    node=OxmlElement("w:tblHeader"); node.set(qn("w:val"),"true"); row._tr.get_or_add_trPr().append(node)


def geometry(section, landscape=False) -> None:
    section.orientation=WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width=Inches(11.7 if landscape else 8.3)
    section.page_height=Inches(8.3 if landscape else 11.7)
    margin=Inches(0.55 if landscape else 0.85)
    section.top_margin=section.bottom_margin=section.left_margin=section.right_margin=margin


def page_field(p) -> None:
    p.add_run("Page "); run=p.add_run()
    begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin")
    instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=" PAGE "
    end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end")
    run._r.extend([begin,instr,end])


def headers_footers(doc: Document) -> None:
    for section in doc.sections:
        section.header.is_linked_to_previous=False
        p=section.header.paragraphs[0]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run("INTERNAL WORKING MANUSCRIPT — NOT FOR SUBMISSION"); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10); r.font.color.rgb=RGBColor(201,76,76)
        section.footer.is_linked_to_previous=False
        p=section.footer.paragraphs[0]; p.clear(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        r=p.add_run("NSS-ID internal evidence package • "); r.font.name="Arial"; r.font.size=Pt(8); r.font.color.rgb=RGBColor(82,102,122); page_field(p)


def styles(doc: Document) -> None:
    s=doc.styles["Normal"]; s.font.name="Times New Roman"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Times New Roman"); s.font.size=Pt(11); s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.15
    for name,size,color in (("Title",22,"C94C4C"),("Heading 1",17,"28527A"),("Heading 2",14,"28527A"),("Heading 3",12,"28527A")):
        s=doc.styles[name]; s.font.name="Arial"; s._element.rPr.rFonts.set(qn("w:eastAsia"),"Arial"); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.keep_with_next=True
    s=doc.styles["Caption"]; s.font.name="Arial"; s.font.size=Pt(9); s.font.color.rgb=RGBColor(82,102,122); s.paragraph_format.keep_with_next=True


def table_captions() -> dict[int,str]:
    out={}; current=None
    for line in CAPTIONS.read_text(encoding="utf-8").splitlines():
        m=re.match(r"## Table S?(\d+)\.",line)
        if m: current=int(m.group(1))
        elif current and line.startswith("**Caption:**"):
            out[current]=plain(line.split("**Caption:**",1)[1]); current=None
    if set(out)!=set(TABLE_FILES): raise ValueError(f"incomplete table captions: {set(out)}")
    return out


def figure_captions(text: str) -> dict[int,str]:
    out={}
    for line in text.splitlines():
        m=re.match(r"\*\*Figure (\d+)\. (.+?)\*\*\s*(.*)",line)
        if m: out[int(m.group(1))]=f"Figure {m.group(1)}. {m.group(2)} {m.group(3)}".strip()
    return out


def csv_rows(path: Path) -> tuple[list[str],list[dict[str,str]]]:
    with path.open(newline="",encoding="utf-8-sig") as f:
        reader=csv.DictReader(f); rows=list(reader); return list(reader.fieldnames or []),rows


def pretty(value: str) -> str:
    special={"wer_percent":"WER (%)","cer_percent":"CER (%)","duration_hours":"Duration (h)","duration_sec":"Duration (s)","mean_duration_sec":"Mean duration (s)","model_family":"Model family","category_english":"Category","percent_of_release_target":"% of release target","path_or_archive_member":"Path / archive member","checksum_or_version":"Checksum / version"}
    return special.get(value,value.replace("_"," ").title())

def add_word_table(doc: Document, number: int, caption: str) -> None:
    fields, rows=csv_rows(TABLE_DIR/TABLE_FILES[number])
    label=f"Table S{number}" if number==6 else f"Table {number}"
    p=doc.add_paragraph(style="Caption"); p.add_run(f"{label}. {caption}").bold=True
    table=doc.add_table(rows=1,cols=len(fields)); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=True
    repeat_header(table.rows[0]); font_size=Pt(6 if len(fields)>=7 else 7 if len(fields)>=5 else 8)
    for i,field in enumerate(fields):
        cell=table.rows[0].cells[i]; cell.text=pretty(field); shade(cell,"28527A"); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.name="Arial"; run.font.size=font_size
    for row in rows:
        cells=table.add_row().cells
        for i,field in enumerate(fields):
            cells[i].text=plain(row.get(field,"")); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows)%2==0: shade(cells[i],"E7F0F8")
            for p2 in cells[i].paragraphs:
                p2.paragraph_format.space_after=Pt(0)
                for run in p2.runs: run.font.name="Arial"; run.font.size=font_size
    p=doc.add_paragraph(style="Caption"); p.add_run(f"Editable source: tables/{TABLE_FILES[number]}")


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    matrix=[]
    for line in lines:
        cells=[plain(x.strip()) for x in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?",x or "") for x in cells): continue
        matrix.append(cells)
    if not matrix: return
    width=max(map(len,matrix)); table=doc.add_table(rows=1,cols=width); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    for i,value in enumerate(matrix[0]):
        cell=table.rows[0].cells[i]; cell.text=value; shade(cell,"28527A")
        for run in cell.paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255); run.font.name="Arial"; run.font.size=Pt(8)
    for values in matrix[1:]:
        cells=table.add_row().cells
        for i in range(width):
            cells[i].text=values[i] if i<len(values) else ""
            if len(table.rows)%2==0: shade(cells[i],"E7F0F8")
            for run in cells[i].paragraphs[0].runs: run.font.name="Arial"; run.font.size=Pt(8)


def insert_figure(doc: Document, number: int, captions: dict[int,str]) -> None:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURE_DIR/FIGURE_FILES[number]),width=Inches(6.5))
    p=doc.add_paragraph(style="Caption"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(plain(captions.get(number,f"Figure {number}."))).bold=True


def build_docx(path: Path) -> None:
    text=MANUSCRIPT.read_text(encoding="utf-8"); tcaps=table_captions(); fcaps=figure_captions(text)
    doc=Document(); styles(doc); geometry(doc.sections[0],False)
    props=doc.core_properties; props.title="NSS-ID internal working manuscript"; props.subject="INTERNAL — NOT FOR SUBMISSION OR PUBLIC RELEASE"; props.author=""; props.last_modified_by=""; props.keywords="internal; not for submission; not for public release"; props.comments="Unresolved material gates remain. Human authorization is required before submission."; props.category="Internal evidence-led working draft"; props.created=FIXED_DATE; props.modified=FIXED_DATE; props.revision=1
    lines=text.splitlines(); inserted=set(); skip_captions=False; i=0
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith("# Figure captions for planned artwork"):
            skip_captions=True; i+=1; continue
        if skip_captions:
            if line.startswith("# References"): skip_captions=False
            else: i+=1; continue
        if line.startswith("|"):
            block=[]
            while i<len(lines) and lines[i].strip().startswith("|"): block.append(lines[i]); i+=1
            add_markdown_table(doc,block); continue
        if not line.strip(): i+=1; continue
        fm=re.search(r"\[FIGURE (\d+) PLACEHOLDER",line)
        if fm:
            number=int(fm.group(1))
            if number in FIGURE_FILES: insert_figure(doc,number,fcaps)
            else:
                p=doc.add_paragraph(); p.style=doc.styles["Caption"]; r=p.add_run(plain(line)); r.bold=True; r.font.color.rgb=RGBColor(201,76,76)
            i+=1; continue
        hm=re.match(r"^(#{1,3})\s+(.+)",line)
        if hm:
            level=len(hm.group(1)); title=plain(hm.group(2))
            if i==0 or "NOT FOR SUBMISSION" in title:
                p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title)
            else: doc.add_heading(title,level=level)
            i+=1; continue
        if line.startswith("- "):
            p=doc.add_paragraph(style="List Bullet"); add_inline(p,line[2:])
        elif re.match(r"^\d+\.\s",line):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.25); p.paragraph_format.first_line_indent=Inches(-.25); add_inline(p,line)
        else:
            p=doc.add_paragraph(); add_inline(p,line)
        for number,trigger in TABLE_TRIGGERS.items():
            if number not in inserted and trigger in line:
                if number in {1,2,3,4,5}: geometry(doc.add_section(WD_SECTION.NEW_PAGE),True)
                add_word_table(doc,number,tcaps[number]); inserted.add(number)
                if number in {1,2,3,4,5}: geometry(doc.add_section(WD_SECTION.NEW_PAGE),False)
        i+=1
    if inserted!=set(TABLE_FILES): raise ValueError(f"tables not inserted: {set(TABLE_FILES)-inserted}")
    headers_footers(doc); doc.save(path); normalize_zip(path)


def normalize_zip(path: Path) -> None:
    with zipfile.ZipFile(path,"r") as source:
        entries=[]
        for info in source.infolist():
            data=source.read(info.filename)
            if info.filename=="docProps/core.xml":
                data=re.sub(rb"(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)",rb"\g<1>2026-07-22T00:00:00Z\g<2>",data)
            entries.append((info,data))
    with tempfile.NamedTemporaryFile(delete=False,dir=path.parent,suffix=path.suffix) as handle: temp=Path(handle.name)
    try:
        with zipfile.ZipFile(temp,"w",compression=zipfile.ZIP_DEFLATED,compresslevel=9) as target:
            for old,data in sorted(entries,key=lambda item:item[0].filename):
                info=zipfile.ZipInfo(old.filename,(1980,1,1,0,0,0)); info.compress_type=zipfile.ZIP_STORED if old.is_dir() else zipfile.ZIP_DEFLATED; info.external_attr=old.external_attr; info.create_system=0
                target.writestr(info,data)
        temp.replace(path)
    finally:
        if temp.exists(): temp.unlink()


def build_xlsx(path: Path) -> None:
    caps=table_captions(); wb=Workbook(); wb.remove(wb.active)
    wb.properties.creator=""; wb.properties.lastModifiedBy=""; wb.properties.title="NSS-ID editable internal tables"; wb.properties.subject="NOT FOR SUBMISSION OR PUBLIC RELEASE"; wb.properties.description="Editable source tables with unresolved gates"; wb.properties.created=FIXED_DATE; wb.properties.modified=FIXED_DATE
    thin=Side(style="thin",color="D9E2EC")
    for idx,(name,filename) in enumerate(SHEETS):
        fields,rows=csv_rows(TABLE_DIR/filename); ws=wb.create_sheet(name); last=max(1,len(fields))
        ws.merge_cells(start_row=1,start_column=1,end_row=1,end_column=last); ws["A1"]="INTERNAL WORKING TABLE — NOT FOR SUBMISSION OR PUBLIC RELEASE"; ws["A1"].font=Font(name="Arial",bold=True,color="FFFFFF",size=12); ws["A1"].fill=PatternFill("solid",fgColor="C94C4C"); ws["A1"].alignment=Alignment(horizontal="center")
        table_label="Table S6" if idx==6 else f"Table {idx}"
        ws.merge_cells(start_row=2,start_column=1,end_row=2,end_column=last); ws["A2"]=("Specifications Table for the NSS-ID internal draft" if idx==0 else f"{table_label}. {caps[idx]}"); ws["A2"].font=Font(name="Arial",bold=True,color="28527A",size=11); ws["A2"].alignment=Alignment(wrap_text=True)
        ws.merge_cells(start_row=3,start_column=1,end_row=3,end_column=last); ws["A3"]=f"Editable source: tables/{filename}"; ws["A3"].font=Font(name="Arial",italic=True,color="52667A",size=9)
        for col,field in enumerate(fields,1):
            c=ws.cell(5,col,pretty(field)); c.font=Font(name="Arial",bold=True,color="FFFFFF",size=9); c.fill=PatternFill("solid",fgColor="28527A"); c.alignment=Alignment(wrap_text=True,vertical="top"); c.border=Border(left=thin,right=thin,top=thin,bottom=thin)
        for rix,row in enumerate(rows,6):
            for col,field in enumerate(fields,1):
                c=ws.cell(rix,col,plain(row.get(field,""))); c.font=Font(name="Arial",size=9); c.alignment=Alignment(wrap_text=True,vertical="top"); c.border=Border(left=thin,right=thin,top=thin,bottom=thin)
                if rix%2==0: c.fill=PatternFill("solid",fgColor="E7F0F8")
        for col,field in enumerate(fields,1):
            longest=max([len(pretty(field)),*[len(str(row.get(field,""))) for row in rows]]); ws.column_dimensions[get_column_letter(col)].width=min(max(12,longest+2),55)
        ws.freeze_panes="A6"; ws.auto_filter.ref=f"A5:{get_column_letter(last)}{5+len(rows)}"; ws.sheet_view.showGridLines=False; ws.page_setup.orientation="landscape" if len(fields)>4 else "portrait"; ws.oddHeader.center.text="&BNOT FOR SUBMISSION&B"; ws.oddFooter.right.text="Page &P of &N"
    wb.save(path); normalize_zip(path)


def write_readme(path: Path) -> None:
    path.write_text(
        """# NSS-ID internal manuscript package — NOT FOR SUBMISSION

> **STOP:** This package is not authorized for journal submission, repository deposition, DOI registration, distribution, or public release.

## Contents
- [`NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx`](NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.docx): canonical internal manuscript built from the supplied Data in Brief v.19 template, with the seven-row Specifications Table, Tables 1–5, and Figures 1–3.
- [`NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md`](NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md): template-aligned canonical Markdown source.
- [`NSS-ID_EVIDENCE_MASTER_INTERNAL.md`](NSS-ID_EVIDENCE_MASTER_INTERNAL.md): longer evidence master retained for claim provenance.
- [`NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx`](NSS-ID_EDITABLE_TABLES_INTERNAL_NOT_FOR_SUBMISSION.xlsx): consolidated editable workbook including Supplementary Table S6.
- [`tables/`](tables/) and [`figures/`](figures/): CSV and PNG/SVG source artifacts.
- [`evidence/`](evidence/): official-template inventory/design, claim registry, line-cited methods evidence matrix, bilingual author questionnaire, journal/example brief, architecture/display controls, verified references, scoring audit, risk controls, and material-gap register.
- [`PACKAGE_MANIFEST.json`](PACKAGE_MANIFEST.json): file hashes and authorization state.
- [Back to the repository Draft Paper index](../README.md).

## Current decision
Internal rebuild: **GO**. Journal submission: **NO-GO**. Public data release: **NO-GO**. Gates G0–G5 remain NO-GO; G6 is unassessed. Ethics/consent, component rights and licence, privacy governance, acquisition evidence, prior-publication eligibility, frozen manifests/checksums, public-or-approved-controlled access, DOI, and all-author authorization remain unresolved.

The benchmark table uses one uniform diagnostic rescore. Historical run-native rankings are provenance only, not publication-comparable results. The complete nine-row table is supplementary by default until per-recipe method cards and sensitivity/interpretation gates close.
""",
        encoding="utf-8",
    )


MARKDOWN_LINK_RE=re.compile(r"(?<!!)\[([^\]]+)\]\(([^)\s]+)\)")
VIRTUAL_PACKAGE=ROOT/"Draft_Paper/05_Submission_Package"


def rewrite_relative_markdown_links(text: str, source: Path, virtual_destination: Path) -> str:
    """Keep copied Markdown links valid from the package's GitHub location."""
    def replace(match: re.Match[str]) -> str:
        label,url=match.group(1),match.group(2)
        if url.startswith(("http://","https://","mailto:","#")): return match.group(0)
        split=urlsplit(url)
        if not split.path: return match.group(0)
        target=(source.parent/unquote(split.path)).resolve()
        try: target.relative_to(ROOT.resolve())
        except ValueError: return match.group(0)
        if not target.exists(): return match.group(0)
        relative=Path(os.path.relpath(target,start=virtual_destination.parent)).as_posix()
        rewritten=quote(relative,safe="/._-")
        if split.fragment: rewritten+=f"#{split.fragment}"
        return f"[{label}]({rewritten})"
    return MARKDOWN_LINK_RE.sub(replace,text)


def copy_for_package(source: Path, destination: Path, virtual_destination: Path) -> None:
    if source.suffix.lower()==".md":
        text=rewrite_relative_markdown_links(source.read_text(encoding="utf-8"),source,virtual_destination)
        destination.write_text(text,encoding="utf-8")
    else: shutil.copyfile(source,destination)


def copy_supporting_files(output: Path) -> None:
    for folder in (output/"tables",output/"figures",output/"evidence"): folder.mkdir(parents=True,exist_ok=True)
    (output/"tables"/"Table_6_frozen_benchmark_validation.csv").unlink(missing_ok=True)
    for _,filename in SHEETS: shutil.copyfile(TABLE_DIR/filename,output/"tables"/filename)
    for source in sorted(FIGURE_DIR.iterdir()):
        if source.is_file(): shutil.copyfile(source,output/"figures"/source.name)
    for source,name in EVIDENCE_FILES:
        if not source.is_file(): raise FileNotFoundError(source)
        copy_for_package(source,output/"evidence"/name,VIRTUAL_PACKAGE/"evidence"/name)
    copy_for_package(TEMPLATE_MANUSCRIPT,output/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md",VIRTUAL_PACKAGE/"NSS-ID_INTERNAL_NOT_FOR_SUBMISSION.md")
    copy_for_package(EVIDENCE_MASTER,output/"NSS-ID_EVIDENCE_MASTER_INTERNAL.md",VIRTUAL_PACKAGE/"NSS-ID_EVIDENCE_MASTER_INTERNAL.md")
    copy_for_package(CAPTIONS,output/"TABLE_CAPTIONS_AND_NOTES_INTERNAL.md",VIRTUAL_PACKAGE/"TABLE_CAPTIONS_AND_NOTES_INTERNAL.md")


def write_manifest(output: Path) -> None:
    files=[]
    for path in sorted(output.rglob("*")):
        if path.is_file() and path.name!="PACKAGE_MANIFEST.json": files.append({"path":path.relative_to(output).as_posix(),"sha256":sha256(path),"bytes":path.stat().st_size})
    source_docx=ROOT/"Draft_Paper/00_Source/Draft Jurnal Data In Brief NSS-ID_ver3.docx"
    if source_docx.exists() and sha256(source_docx)!=SOURCE_DOCX_SHA256:
        raise ValueError("local source DOCX hash does not match the recorded custody hash")
    payload={"status":"internal_not_for_submission","submission_authorized":False,"public_release_authorized":False,"gate_status":{"G0":"NO-GO","G1":"NO-GO","G2":"NO-GO","G3":"NO-GO","G4":"NO-GO","G5":"NO-GO","G6":"UNASSESSED"},"source_docx_sha256":SOURCE_DOCX_SHA256,"official_template":"Draft_Paper/data-in-brief-article-template.docx","official_template_version":"v.19 (December 2024)","official_template_sha256":sha256(dib_template.TEMPLATE),"template_manuscript_source":"Draft_Paper/04_Revised_Draft/06_DATA_IN_BRIEF_TEMPLATE_MANUSCRIPT.md","template_manuscript_source_sha256":sha256(TEMPLATE_MANUSCRIPT),"template_generator":"Draft_Paper/99_Admin/build_data_in_brief_template_docx.py","template_generator_sha256":sha256(Path(dib_template.__file__)),"generator":"Draft_Paper/99_Admin/build_internal_docx_package.py","generator_sha256":sha256(Path(__file__)),"files":files}
    (output/"PACKAGE_MANIFEST.json").write_text(json.dumps(payload,indent=2)+"\n",encoding="utf-8")


def build(output: Path) -> None:
    output.mkdir(parents=True,exist_ok=True); copy_supporting_files(output); dib_template.build(output/DOCX_NAME); build_xlsx(output/XLSX_NAME); write_readme(output/"README_INTERNAL_NOT_FOR_SUBMISSION.md"); write_manifest(output)
    print(f"Built internal-only package at {output}")


def main() -> None:
    p=argparse.ArgumentParser(description=__doc__); p.add_argument("--output-dir",type=Path,default=DEFAULT_OUTPUT); args=p.parse_args(); build(args.output_dir.resolve())


if __name__=="__main__": main()
