#!/usr/bin/env python3
"""Regression tests for the official Data in Brief v.19 template-based DOCX."""

from __future__ import annotations

import hashlib
import re
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = ROOT / "Draft_Paper" / "data-in-brief-article-template.docx"
SCRIPT = ROOT / "Draft_Paper" / "99_Admin" / "build_data_in_brief_template_docx.py"
SOURCE = ROOT / "Draft_Paper" / "04_Revised_Draft" / "06_DATA_IN_BRIEF_TEMPLATE_MANUSCRIPT.md"
LEDGER = ROOT / "Draft_Paper" / "04_Revised_Draft" / "03_MATERIAL_GAP_PLACEHOLDERS.md"
TEMPLATE_SHA256 = "5c02d5f9e0762e05f69c06d1d042ea800b6214427c82a78166863dfd17264190"
OFFICIAL_HEADINGS = [
    "ARTICLE INFORMATION",
    "SPECIFICATIONS TABLE",
    "VALUE OF THE DATA",
    "BACKGROUND",
    "DATA DESCRIPTION",
    "EXPERIMENTAL DESIGN, MATERIALS AND METHODS",
    "LIMITATIONS",
    "ETHICS STATEMENT",
    "CRediT AUTHOR STATEMENT",
    "ACKNOWLEDGEMENTS",
    "DECLARATION OF COMPETING INTERESTS",
    "REFERENCES",
]
SPECIFICATION_LABELS = [
    "Subject",
    "Specific subject area",
    "Type of data",
    "Data collection",
    "Data source location",
    "Data accessibility",
    "Related research article",
]


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def document_text(doc: Document) -> str:
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    chunks.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return "\n".join(chunks)


def section_text(doc: Document, heading: str, next_heading: str) -> str:
    paragraphs = doc.paragraphs
    start = next(i for i, paragraph in enumerate(paragraphs) if paragraph.text.strip() == heading)
    end = next(i for i, paragraph in enumerate(paragraphs[start + 1 :], start + 1) if paragraph.text.strip() == next_heading)
    return "\n".join(paragraph.text for paragraph in paragraphs[start + 1 : end])


class DataInBriefTemplateDocxTests(unittest.TestCase):
    def build(self, output: Path) -> None:
        subprocess.run(
            [sys.executable, str(SCRIPT), "--output", str(output)],
            cwd=ROOT,
            check=True,
            timeout=300,
        )

    def test_template_based_docx_is_conformant_sanitized_and_deterministic(self) -> None:
        self.assertEqual(TEMPLATE_SHA256, sha256(TEMPLATE))
        self.assertTrue(SOURCE.is_file())
        with tempfile.TemporaryDirectory() as first_dir, tempfile.TemporaryDirectory() as second_dir:
            first = Path(first_dir) / "nssid-dib.docx"
            second = Path(second_dir) / "nssid-dib.docx"
            self.build(first)
            self.build(second)
            self.assertEqual(sha256(first), sha256(second))

            doc = Document(first)
            text = document_text(doc)
            headings = [p.text.strip() for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
            self.assertEqual(OFFICIAL_HEADINGS, headings)
            self.assertNotIn("AUTHOR INSTRUCTIONS", text)
            self.assertNotIn("TECHNICAL VALIDATION", headings)
            self.assertNotIn("DATA AVAILABILITY", headings)
            self.assertNotIn("FUNDING", headings)
            self.assertIn("INTERNAL WORKING DRAFT — NOT FOR SUBMISSION OR PUBLIC RELEASE", text)

            title_line = next(p.text for p in doc.paragraphs if p.text.startswith("Article title: "))
            title = title_line.split(":", 1)[1].strip()
            self.assertRegex(title.lower(), r"\b(data|dataset)\b")
            keywords_line = next(p.text for p in doc.paragraphs if p.text.startswith("Keywords: "))
            keywords = [item.strip() for item in keywords_line.split(":", 1)[1].split(";") if item.strip()]
            self.assertGreaterEqual(len(keywords), 4)
            self.assertLessEqual(len(keywords), 8)
            title_words = set(re.findall(r"[a-z0-9]+", title.lower()))
            for keyword in keywords:
                self.assertTrue(title_words.isdisjoint(re.findall(r"[a-z0-9]+", keyword.lower())), keyword)

            abstract_line = next(p.text for p in doc.paragraphs if p.text.startswith("Abstract: "))
            abstract = abstract_line.split(":", 1)[1].strip()
            abstract_words = re.findall(r"\b[\w'-]+\b", abstract)
            self.assertGreaterEqual(len(abstract_words), 100)
            self.assertLessEqual(len(abstract_words), 500)
            for phrase in ["therefore", "demonstrates that", "shows that", "we conclude"]:
                self.assertNotIn(phrase, abstract.lower())

            background = section_text(doc, "BACKGROUND", "DATA DESCRIPTION")
            limitations = section_text(doc, "LIMITATIONS", "ETHICS STATEMENT")
            self.assertLessEqual(len(re.findall(r"\b[\w'-]+\b", background)), 200)
            self.assertLessEqual(len(re.findall(r"\b[\w'-]+\b", limitations)), 200)
            value_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "VALUE OF THE DATA")
            background_start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "BACKGROUND")
            bullets = [p for p in doc.paragraphs[value_start + 1 : background_start] if "List" in (p.style.name if p.style else "")]
            self.assertGreaterEqual(len(bullets), 3)
            self.assertLessEqual(len(bullets), 6)
            self.assertTrue(all(len(re.findall(r"\b[\w'-]+\b", p.text)) <= 150 for p in bullets))

            self.assertEqual(6, len(doc.tables))
            specifications = doc.tables[0]
            self.assertEqual(SPECIFICATION_LABELS, [row.cells[0].text.strip() for row in specifications.rows])
            specific_subject = re.sub(r"\s+", "", specifications.rows[1].cells[1].text)
            data_collection = re.sub(r"\s+", "", specifications.rows[3].cells[1].text)
            self.assertLessEqual(len(specific_subject), 150)
            self.assertLessEqual(len(data_collection), 600)
            self.assertNotIn("Internal evidence status", specifications.rows[3].cells[1].text)
            self.assertEqual(3, len(doc.inline_shapes))
            references = section_text(doc, "REFERENCES", "__END__") if False else text.split("REFERENCES", 1)[1]
            numbered = re.findall(r"(?m)^\[(\d+)\]\s", references)
            self.assertLessEqual(len(numbered), 20)
            self.assertIn("10.15587/1729-4061.2026.350949", references.splitlines()[1])
            self.assertIn("[dataset]", references)

            canonical = set(re.findall(r"`(\[MATERIAL GAP:[^`]+\])`", LEDGER.read_text(encoding="utf-8")))
            self.assertEqual(33, len(canonical))
            observed = set(re.findall(r"\[MATERIAL GAP:[^]]+\]", text))
            self.assertEqual(canonical, observed)

            section = doc.sections[0]
            self.assertAlmostEqual(8.27, section.page_width.inches, places=1)
            self.assertAlmostEqual(11.69, section.page_height.inches, places=1)
            for margin in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin]:
                self.assertAlmostEqual(1.0, margin.inches, places=1)
            self.assertEqual(16.0, doc.styles["Heading 1"].font.size.pt)

            with zipfile.ZipFile(first) as archive, zipfile.ZipFile(TEMPLATE) as template_archive:
                names = set(archive.namelist())
                self.assertNotIn("word/comments.xml", names)
                self.assertNotIn("word/commentsExtended.xml", names)
                self.assertNotIn("word/commentsIds.xml", names)
                self.assertNotIn("word/commentsExtensible.xml", names)
                self.assertNotIn("docProps/custom.xml", names)
                self.assertTrue(all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist()))
                xml = "\n".join(
                    archive.read(name).decode("utf-8", "ignore")
                    for name in names
                    if name.endswith((".xml", ".rels"))
                )
                self.assertNotIn("documentProtection", xml)
                self.assertNotIn("commentRangeStart", xml)
                self.assertNotIn("TargetMode=\"External\"", xml)
                self.assertNotIn("Please fill in the template below", xml)
                header_images = [name for name in names if name.startswith("word/media/")]
                expected_header_hash = hashlib.sha256(template_archive.read("word/media/image1.png")).hexdigest()
                self.assertIn(expected_header_hash, {hashlib.sha256(archive.read(name)).hexdigest() for name in header_images})


if __name__ == "__main__":
    unittest.main()
